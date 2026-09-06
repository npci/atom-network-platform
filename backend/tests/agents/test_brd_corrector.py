# Copyright 2026 National Payments Corporation of India
# SPDX-License-Identifier: MIT

"""brd_corrector — derive verbatim-locatable BRD corrections (LLM) + apply them to
the uploaded .docx as a new BRD version (formatting/images preserved)."""
import asyncio
import json
from types import SimpleNamespace

from docx import Document

import app.agents.brd_corrector as BC
from app.agents.brd_corrector import propose_brd_corrections, apply_doc_corrections


class _Q:
    def __init__(self, r): self._r = r
    def filter(self, *a, **k): return self
    def order_by(self, *a, **k): return self
    def first(self): return self._r


class _DB:
    def __init__(self, brd):
        self._brd = brd
        self.added = []
        self.committed = False
    def query(self, *a, **k): return _Q(self._brd)
    def add(self, o): self.added.append(o)
    def commit(self): self.committed = True
    def rollback(self): pass


# ── propose (LLM) ────────────────────────────────────────────────────────────
def test_propose_keeps_only_verbatim_locatable(monkeypatch):
    async def _llm(*a, **kw):
        return json.loads('{"corrections":[{"conflict_id":"c1","find":"Rs 5,000 cap","replace":"no cap"},'
                '{"conflict_id":"c2","find":"NOT IN THE BRD","replace":"x"}]}')
    monkeypatch.setattr(BC, "call_llm_structured", _llm)
    out = asyncio.run(propose_brd_corrections("The Rs 5,000 cap applies to P2P.",
                                              [{"id": "c1", "text": "cap"}, {"id": "c2", "text": "y"}]))
    assert len(out) == 1 and out[0]["conflict_id"] == "c1"      # c2 not verbatim → dropped
    assert out[0]["find"] == "Rs 5,000 cap" and out[0]["replace"] == "no cap"


def test_propose_empty_inputs():
    assert asyncio.run(propose_brd_corrections("", [{"id": "c1"}])) == []
    assert asyncio.run(propose_brd_corrections("text", [])) == []


def test_propose_fails_open(monkeypatch):
    async def _boom(**kw): raise RuntimeError("down")
    monkeypatch.setattr(BC, "call_llm_structured", _boom)
    assert asyncio.run(propose_brd_corrections("Rs 5,000", [{"id": "c1", "text": "x"}])) == []


# ── apply (docx → new BRD version) ───────────────────────────────────────────
def test_apply_docx_creates_corrected_version(tmp_path):
    src = tmp_path / "brd.docx"
    d = Document(); d.add_paragraph("The cap is Rs 5,000 for all flows."); d.save(str(src))
    brd = SimpleNamespace(change_request_id="cr", version=1, content="The cap is Rs 5,000 for all flows.",
                          file_path=str(src), original_filename="brd.docx")
    db = _DB(brd)
    v = apply_doc_corrections(db, "cr", "brd",[{"find": "Rs 5,000", "replace": "no cap"}], corrected_by="u")
    assert v == 2 and db.committed and db.added
    new = db.added[0]
    assert new.version == 2 and "no cap" in new.content and new.source.value == "uploaded"


def test_apply_md_creates_corrected_version(tmp_path):
    src = tmp_path / "brd.md"
    src.write_text("# BRD\nThe cap is Rs 5,000.", encoding="utf-8")
    brd = SimpleNamespace(change_request_id="cr", version=1, content="# BRD\nThe cap is Rs 5,000.",
                          file_path=str(src), original_filename="brd.md")
    db = _DB(brd)
    v = apply_doc_corrections(db, "cr", "brd",[{"find": "Rs 5,000", "replace": "no cap"}])
    assert v == 2 and "no cap" in db.added[0].content


def test_apply_noop_when_nothing_matches(tmp_path):
    src = tmp_path / "brd.md"
    src.write_text("nothing to change", encoding="utf-8")
    brd = SimpleNamespace(change_request_id="cr", version=1, content="nothing to change",
                          file_path=str(src), original_filename="b.md")
    db = _DB(brd)
    assert apply_doc_corrections(db, "cr", "brd",[{"find": "absent", "replace": "x"}]) is None
    assert db.added == []


def test_apply_additions_adds_back_omission(tmp_path):
    # Flaw-1 fix: a dropped plan requirement (omission) resolved plan-wins is ADDED
    # back (find/replace can't insert), producing a new version even with no corrections.
    src = tmp_path / "brd.docx"
    d = Document(); d.add_paragraph("Overview only."); d.save(str(src))
    brd = SimpleNamespace(change_request_id="cr", version=1, content="Overview only.",
                          file_path=str(src), original_filename="brd.docx")
    db = _DB(brd)
    v = apply_doc_corrections(db, "cr", "brd",[], additions=["Add enum value SPLIT to payConstant"])
    assert v == 2 and db.added
    assert "SPLIT" in db.added[0].content   # requirement added back to the BRD


def test_apply_surfaces_unmatched_correction(tmp_path):
    # G7: a correction whose find is absent is surfaced as a follow-up note, not dropped.
    src = tmp_path / "brd.md"
    src.write_text("The cap is Rs 5,000.", encoding="utf-8")
    brd = SimpleNamespace(change_request_id="cr", version=1, content="The cap is Rs 5,000.",
                          file_path=str(src), original_filename="brd.md")
    db = _DB(brd)
    v = apply_doc_corrections(db, "cr", "brd",
                              [{"find": "Rs 5,000", "replace": "no cap"},
                               {"find": "TOTALLY ABSENT PHRASE", "replace": "x"}])
    assert v == 2
    content = db.added[0].content
    assert "no cap" in content
    assert "TOTALLY ABSENT PHRASE" in content and "manual edit" in content.lower()


def test_apply_tech_spec_doc_kind(tmp_path):
    # G4: the corrector works for tech_spec, not just brd.
    src = tmp_path / "tsd.md"
    src.write_text("The cap is Rs 5,000.", encoding="utf-8")
    tsd = SimpleNamespace(change_request_id="cr", version=1, content="The cap is Rs 5,000.",
                          file_path=str(src), original_filename="tsd.md")
    db = _DB(tsd)
    v = apply_doc_corrections(db, "cr", "tech_spec", [{"find": "Rs 5,000", "replace": "no cap"}])
    assert v == 2 and "no cap" in db.added[0].content and db.added[0].source.value == "uploaded"


def test_propose_uses_custom_target(monkeypatch):
    # G5: a custom resolution's text becomes the TARGET the corrector edits toward.
    seen = {}
    async def _llm(*a, **kw):
        seen["content"] = a[1]
        return json.loads('{"corrections":[{"conflict_id":"c1","find":"Rs 5,000 cap","replace":"cap only for P2P"}]}')
    monkeypatch.setattr(BC, "call_llm_structured", _llm)
    out = asyncio.run(propose_brd_corrections("The Rs 5,000 cap applies.",
                                              [{"id": "c1", "text": "cap", "_target": "cap only for P2P"}]))
    assert out and out[0]["replace"] == "cap only for P2P"
    assert "cap only for P2P" in seen["content"]   # the custom TARGET reached the prompt


# ── uncorrected_followups (breadcrumb for ungrounded plan-wins/custom conflicts) ──
def test_uncorrected_followups_flags_ungrounded_conflicts():
    # A plan-wins/custom conflict the corrector couldn't ground into a verbatim edit
    # must become a manual-review follow-up — never a silent no-op.
    from app.agents.brd_corrector import uncorrected_followups
    to_edit = [
        {"id": "c1", "text": "cap of Rs 5,000", "_target": "the ratified plan"},
        {"id": "c2", "text": "reviewer ruling item", "_target": "cap only for P2P"},
    ]
    corrections = [{"conflict_id": "c1", "find": "Rs 5,000", "replace": "no cap"}]  # only c1 grounded
    outs = uncorrected_followups(to_edit, corrections)
    assert len(outs) == 1                                        # c1 corrected; c2 ungrounded → 1 followup
    assert "cap only for P2P" in outs[0] and "reviewer ruling item" in outs[0]
    assert "revise manually" in outs[0].lower()


def test_uncorrected_followups_empty_when_all_corrected():
    from app.agents.brd_corrector import uncorrected_followups
    to_edit = [{"id": "c1", "text": "x"}]
    assert uncorrected_followups(to_edit, [{"conflict_id": "c1", "find": "a", "replace": "b"}]) == []
