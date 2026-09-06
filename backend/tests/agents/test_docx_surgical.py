# Copyright 2026 National Payments Corporation of India
# SPDX-License-Identifier: MIT

"""docx_surgical (Path B) — run-split-aware in-place .docx correction with the
byte-identical image guarantee. Builds a real .docx (text + split runs + table +
embedded image) and asserts text edits land while images survive untouched."""
import io

from docx import Document
from docx.shared import Inches
from PIL import Image

from app.services.docx_surgical import correct_docx, _replace_in_runs, _media_hashes


def _png():
    buf = io.BytesIO()
    Image.new("RGB", (12, 12), (200, 30, 30)).save(buf, format="PNG")
    buf.seek(0)
    return buf


def _make_docx(path):
    doc = Document()
    doc.add_paragraph("The per-transaction cap is Rs 5,000 for all flows.")
    p = doc.add_paragraph()                       # a phrase SPLIT across runs
    p.add_run("The amount is ")
    p.add_run("Rs 5,")
    p.add_run("000")
    p.add_run(" per participant.")
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Cap"
    t.rows[0].cells[1].text = "Rs 5,000"
    doc.add_picture(_png(), width=Inches(1))      # embedded image
    doc.save(path)


def test_replace_in_runs_handles_split_phrase():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("cap "); p.add_run("Rs 5,"); p.add_run("000 x")
    assert _replace_in_runs(p.runs, "Rs 5,000", "no cap") is True
    assert "".join(r.text for r in p.runs) == "cap no cap x"


def test_replace_in_runs_untouched_when_absent():
    doc = Document()
    p = doc.add_paragraph("nothing to change here")
    assert _replace_in_runs(p.runs, "absent", "x") is False
    assert p.text == "nothing to change here"


def test_correct_docx_edits_text_and_preserves_image(tmp_path):
    src, out = str(tmp_path / "in.docx"), str(tmp_path / "out.docx")
    _make_docx(src)
    before_media = _media_hashes(src)
    assert before_media                                   # an image is present

    res = correct_docx(src, [
        {"find": "Rs 5,000 for all flows", "replace": "no cap"},   # single-run
        {"find": "Rs 5,000", "replace": "no cap"},                 # split-run
        {"find": "does-not-exist", "replace": "x"},                # unmatched
    ], out)

    assert res["applied"] == 2
    assert res["unmatched"] == ["does-not-exist"]
    assert res["media_preserved"] is True                 # ← image byte-identical

    d = Document(out)
    text = "\n".join(p.text for p in d.paragraphs)
    assert "no cap" in text and "5,000 for all flows" not in text
    assert _media_hashes(out) == before_media             # images unchanged, provably


def test_correct_docx_empty_is_noop(tmp_path):
    src, out = str(tmp_path / "in.docx"), str(tmp_path / "out.docx")
    _make_docx(src)
    res = correct_docx(src, [], out)
    assert res["applied"] == 0 and res["media_preserved"] is True
