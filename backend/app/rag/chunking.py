# Copyright 2026 National Payments Corporation of India
# SPDX-License-Identifier: MIT

"""Structure-preserving chunker for the knowledge base.

Each supported file type has a dedicated chunker that preserves semantic
boundaries (section headers for PDF/DOCX, rows for XLSX, elements for XSD).
Every chunk carries metadata so downstream retrieval can cite precisely.

Returns list[dict] with:
    content:       str  (the chunk text)
    section_title: str | None
    page:          int | None
    chunk_type:    "text" | "table" | "row" | "element"

Scanned PDFs (no text layer) automatically fall back to OCR via Tesseract
when pytesseract + pypdfium2 are installed.
"""
import logging
import re
import shutil
from pathlib import Path

logger = logging.getLogger(__name__)

# Windows Tesseract default install location (used if tesseract not on PATH)
_WINDOWS_TESSERACT_FALLBACK = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

MAX_CHUNK_CHARS = 1800
MIN_CHUNK_CHARS = 80
DEFAULT_CHUNK_SIZE = 800
DEFAULT_CHUNK_OVERLAP = 100

_SECTION_RE = re.compile(
    r"^\s*(\d+(?:\.\d+){0,3})\s+([A-Z][^\n]{3,120})\s*$",
    re.M,
)

# langchain 1.x removed `langchain.text_splitter`; the splitters live in their
# own distribution now (pinned in requirements.txt as langchain-text-splitters).
# Imported lazily inside _split_section() because langchain-text-splitters 1.x
# imports sentence_transformers at its __init__.py level, which pulls in torch
# and triggers a torch._dynamo crash ("Duplicate dispatch rule for sys.intern")
# on certain Python 3.11 builds at import time. Lazy import defers the entire
# chain until chunking is actually needed — startup stays clean.
_fallback_splitter = None


def _get_fallback_splitter():
    """Lazy-init the RecursiveCharacterTextSplitter.

    Imported lazily because langchain-text-splitters 1.x pulls in
    sentence_transformers → torch at its __init__.py level, which crashes
    certain Python 3.11 builds with a torch._dynamo dispatch error.
    """
    global _fallback_splitter
    if _fallback_splitter is None:
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        _fallback_splitter = RecursiveCharacterTextSplitter(
            chunk_size=DEFAULT_CHUNK_SIZE,
            chunk_overlap=DEFAULT_CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
    return _fallback_splitter


def _split_section(text: str, section_title: str | None, page: int | None) -> list[dict]:
    text = text.strip()
    if len(text) < MIN_CHUNK_CHARS:
        return []
    if len(text) <= MAX_CHUNK_CHARS:
        return [{
            "content": text,
            "section_title": section_title,
            "page": page,
            "chunk_type": "text",
        }]
    splitter = _get_fallback_splitter()
    return [
        {"content": p.strip(), "section_title": section_title, "page": page, "chunk_type": "text"}
        for p in splitter.split_text(text)
        if len(p.strip()) >= MIN_CHUNK_CHARS
    ]


# ── OCR fallback ─────────────────────────────────────────────────────────────

def _ocr_pdf(path: Path) -> list[tuple[int, str]]:
    """Run Tesseract OCR on a scanned PDF. Returns list of (page_num, text).

    Called only when PyMuPDF text extraction yields empty content. Requires
    pytesseract + pypdfium2 + Tesseract binary; returns [] gracefully if
    any piece is missing.
    """
    try:
        import pytesseract
        import pypdfium2 as pdfium
    except ImportError:
        logger.warning("OCR deps missing (pytesseract/pypdfium2). Skipping OCR: %s", path.name)
        return []

    # Locate tesseract binary (PATH first, then Windows default)
    if not shutil.which("tesseract") and Path(_WINDOWS_TESSERACT_FALLBACK).exists():
        pytesseract.pytesseract.tesseract_cmd = _WINDOWS_TESSERACT_FALLBACK

    try:
        pdf = pdfium.PdfDocument(str(path))
    except Exception as e:
        logger.error("OCR: cannot open PDF %s — %s", path, e)
        return []

    pages_text: list[tuple[int, str]] = []
    try:
        for pno in range(len(pdf)):
            page = pdf[pno]
            img = page.render(scale=2).to_pil()
            try:
                text = pytesseract.image_to_string(img, lang="eng")
            except pytesseract.TesseractNotFoundError:
                logger.error("OCR: Tesseract binary not found. Install from https://github.com/UB-Mannheim/tesseract/wiki")
                return []
            except Exception as e:
                logger.warning("OCR: page %d of %s failed: %s", pno + 1, path.name, e)
                continue
            if text.strip():
                pages_text.append((pno + 1, text))
    finally:
        pdf.close()

    logger.info("OCR extracted %d pages from %s", len(pages_text), path.name)
    return pages_text


# ── PDF ──────────────────────────────────────────────────────────────────────

def _extract_pdf_pages(path: Path) -> list[tuple[int, str]]:
    """Per-page ``(page_no, text)`` from a born-digital PDF's text layer.

    Tries pypdfium2 (PDFium) first — best layout fidelity, and the same engine
    ``_ocr_pdf`` renders with — then falls back to pypdf. PDFium is strict and
    raises "Data format error" on spec-violating PDFs (broken xref, malformed
    objects) that the AGPL pymupdf we replaced used to tolerate; pypdf parses
    non-strict by default and recovers from most of those. Returns ``[]`` only
    when NEITHER backend can open the file (caller then tries OCR)."""
    # 1) pypdfium2 — preferred.
    try:
        import pypdfium2 as pdfium
        doc = pdfium.PdfDocument(str(path))
        try:
            return [(pno + 1, doc[pno].get_textpage().get_text_range() or "")
                    for pno in range(len(doc))]
        finally:
            doc.close()
    except Exception as e:
        logger.warning("pypdfium2 could not read %s (%s) — falling back to pypdf", path.name, e)

    # 2) pypdf — lenient pure-Python parser for files PDFium rejects.
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))  # strict=False by default → recovers broken structure
        return [(i + 1, (pg.extract_text() or "")) for i, pg in enumerate(reader.pages)]
    except Exception as e:
        logger.error("PDF open failed (pypdfium2 + pypdf both failed): %s — %s", path, e)
        return []


def chunk_pdf_sections(path: Path) -> list[dict]:
    pages = _extract_pdf_pages(path)

    # Scanned PDF (or a file no text backend could open) → OCR the whole doc.
    if not any(text.strip() for _, text in pages):
        logger.info("No text layer in %s — running OCR (slow)", path.name)
        pages = _ocr_pdf(path)
    if not pages:
        return []

    full_text_parts: list[str] = []
    page_offsets: list[tuple[int, int]] = []
    offset = 0
    for pno, text in pages:
        page_offsets.append((offset, pno))
        full_text_parts.append(text)
        offset += len(text) + 1
    full_text = "\n".join(full_text_parts)

    def _page_for_offset(idx: int) -> int | None:
        last = None
        for off, pno in page_offsets:
            if off <= idx:
                last = pno
            else:
                break
        return last

    matches = list(_SECTION_RE.finditer(full_text))
    if not matches:
        # Unstructured PDF — fall back to recursive splitter over full text
        return _split_section(full_text, None, None)

    chunks: list[dict] = []

    # Preamble before first section header
    preamble = full_text[: matches[0].start()].strip()
    if preamble:
        chunks.extend(_split_section(preamble, None, _page_for_offset(0)))

    for i, m in enumerate(matches):
        end = matches[i + 1].start() if i + 1 < len(matches) else len(full_text)
        section_title = f"{m.group(1)} {m.group(2).strip()}"
        body = full_text[m.end():end]
        chunks.extend(_split_section(body, section_title, _page_for_offset(m.start())))

    return chunks


# ── DOCX ─────────────────────────────────────────────────────────────────────

def chunk_docx_sections(path: Path) -> list[dict]:
    from docx import Document
    try:
        doc = Document(str(path))
    except Exception as e:
        logger.error("DOCX open failed: %s — %s", path, e)
        return []

    chunks: list[dict] = []
    current_section: str | None = None
    buffer: list[str] = []

    def _flush():
        nonlocal buffer
        if buffer:
            chunks.extend(_split_section("\n".join(buffer), current_section, None))
            buffer = []

    for p in doc.paragraphs:
        style = (p.style.name or "") if p.style else ""
        text = p.text.strip()
        if not text:
            continue
        if style.startswith("Heading"):
            _flush()
            current_section = text
        else:
            buffer.append(text)
    _flush()

    for t_idx, tbl in enumerate(doc.tables):
        rows = [" | ".join(c.text.strip() for c in row.cells) for row in tbl.rows]
        content = "\n".join(rows).strip()
        if len(content) >= MIN_CHUNK_CHARS:
            label = f"Table {t_idx + 1}"
            if current_section:
                label += f" ({current_section})"
            chunks.append({
                "content": content,
                "section_title": label,
                "page": None,
                "chunk_type": "table",
            })

    return chunks


# ── XLSX ─────────────────────────────────────────────────────────────────────

def chunk_xlsx_rows(path: Path) -> list[dict]:
    import openpyxl
    try:
        wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    except Exception as e:
        logger.error("XLSX open failed: %s — %s", path, e)
        return []

    chunks: list[dict] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(c).strip() if c is not None else "" for c in rows[0]]
        for r_idx, row in enumerate(rows[1:], start=2):
            if all(c is None or str(c).strip() == "" for c in row):
                continue
            parts = []
            for h, v in zip(headers, row):
                if v is None:
                    continue
                val = str(v).strip()
                if not val:
                    continue
                parts.append(f"{h}: {val}" if h else val)
            content = " | ".join(parts)
            if len(content) < MIN_CHUNK_CHARS:
                continue
            chunks.append({
                "content": content,
                "section_title": f"{sheet_name} — row {r_idx}",
                "page": None,
                "chunk_type": "row",
            })
    wb.close()
    return chunks


# ── XML / XSD ────────────────────────────────────────────────────────────────

def chunk_xml_elements(path: Path) -> list[dict]:
    from lxml import etree
    # Explicit parser, DEFENCE IN DEPTH — not a fix for a live hole. The input is
    # an ingested document (whatever a user uploaded), so it is untrusted, and an
    # external SAST report flagged this call as XXE. Measured on the pinned lxml
    # (6.1.0) and the previous pin (5.3.0), the bare call is NOT exploitable:
    # external SYSTEM entities are refused ("Entity not defined") and libxml2
    # caps entity amplification, so neither file disclosure nor a billion-laughs
    # bomb gets through. What the default DOES do is expand internal entities.
    #
    # Setting it anyway costs nothing, states the intent locally, survives a
    # future change to lxml/libxml2 defaults, and matches the four sibling XSD
    # parsers (xsd_context, xsd_graph_builder, jaxb_mapper,
    # module_context_generator) which all pass exactly this parser.
    parser = etree.XMLParser(resolve_entities=False, no_network=True)
    try:
        tree = etree.parse(str(path), parser)
    except Exception as e:
        logger.error("XML parse failed: %s — %s", path, e)
        return []

    root = tree.getroot()
    ns = {"xs": "http://www.w3.org/2001/XMLSchema"}
    targets = root.findall(".//xs:element", ns) + root.findall(".//xs:complexType", ns)

    if targets:
        chunks: list[dict] = []
        for el in targets:
            name = el.get("name") or "anonymous"
            xml_str = etree.tostring(el, pretty_print=True, encoding="unicode")
            if len(xml_str.strip()) < MIN_CHUNK_CHARS:
                continue
            tag = el.tag.split("}")[-1]
            chunks.append({
                "content": xml_str,
                "section_title": f"{tag}: {name}",
                "page": None,
                "chunk_type": "element",
            })
        return chunks

    xml_str = etree.tostring(root, pretty_print=True, encoding="unicode")
    return _split_section(xml_str, None, None)


# ── Plain text ───────────────────────────────────────────────────────────────

def chunk_text(path: Path) -> list[dict]:
    try:
        raw = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        logger.error("Text read failed: %s — %s", path, e)
        return []
    return _split_section(raw, None, None)


# ── Dispatcher ───────────────────────────────────────────────────────────────

_DISPATCH = {
    ".pdf":  chunk_pdf_sections,
    ".docx": chunk_docx_sections,
    ".xlsx": chunk_xlsx_rows,
    ".xml":  chunk_xml_elements,
    ".xsd":  chunk_xml_elements,
    ".txt":  chunk_text,
    ".md":   chunk_text,
}

SUPPORTED_EXTENSIONS = set(_DISPATCH.keys())


def chunk_file(path: Path) -> list[dict]:
    """Return structured chunks for a supported file. [] on failure/unsupported."""
    ext = path.suffix.lower()
    fn = _DISPATCH.get(ext)
    if fn is None:
        logger.warning("Unsupported file type: %s", path)
        return []
    try:
        return fn(path)
    except Exception as e:
        logger.error("Chunking failed: %s — %s", path, e)
        return []
