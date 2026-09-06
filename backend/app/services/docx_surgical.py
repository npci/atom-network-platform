# Copyright 2026 National Payments Corporation of India
# SPDX-License-Identifier: MIT

"""Surgical in-place editing of a user-uploaded .docx (reconciliation Path B).

Applies targeted text corrections to an uploaded BRD WITHOUT re-rendering it, so
the user's formatting, tables and — critically — embedded images survive
unchanged. Uploaded docs have no docgen block-IDs; the .docx XML tree IS the
addressing system. We locate text by content, handle Word's run-splitting (a
phrase fragmented across <w:r> runs), and put the replacement in the first
matched run so its formatting is preserved. Runs the match doesn't touch are
left as-is.

Image guarantee: every word/media/* part is checksummed before/after — a
``media_preserved=False`` result means an image changed and the caller MUST
discard the output (the byte-identical-images promise was violated).
"""
from __future__ import annotations

import hashlib
import zipfile


def _media_hashes(path: str) -> dict:
    """sha256 of every word/media/* part — the image-preservation fingerprint."""
    out: dict = {}
    with zipfile.ZipFile(path) as z:
        for n in z.namelist():
            if n.startswith("word/media/"):
                out[n] = hashlib.sha256(z.read(n)).hexdigest()
    return out


def _iter_paragraphs(container):
    """Every paragraph in the body + (recursively) every table cell."""
    for p in container.paragraphs:
        yield p
    for t in container.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from _iter_paragraphs(cell)


def _replace_in_runs(runs, find: str, replace: str) -> bool:
    """Replace the FIRST occurrence of ``find`` across a paragraph's runs, even
    when Word split it across several <w:r>. The replacement lands in the first
    matched run (its formatting is kept); the matched text is removed from the
    rest. Runs the match doesn't overlap are untouched. Returns True if changed."""
    text = "".join(r.text for r in runs)
    idx = text.find(find)
    if idx < 0:
        return False
    end = idx + len(find)
    pos = 0
    new_texts = [r.text for r in runs]
    for i, r in enumerate(runs):
        r_start, r_end = pos, pos + len(r.text)
        pos = r_end
        if r_end <= idx or r_start >= end:
            continue  # run entirely outside the match — leave it alone
        local_start = max(0, idx - r_start)
        local_end = min(len(r.text), end - r_start)
        before, after = r.text[:local_start], r.text[local_end:]
        new_texts[i] = (before + replace + after) if (r_start <= idx < r_end) else (before + after)
    for i, r in enumerate(runs):
        if r.text != new_texts[i]:
            r.text = new_texts[i]
    return True


def correct_docx(src_path: str, corrections: list[dict], out_path: str,
                 additions: list[str] | None = None) -> dict:
    """Apply ``[{find, replace}]`` corrections to ``src_path`` → ``out_path``, and
    optionally APPEND ``additions`` (dropped plan requirements to add back — the
    corrector can't insert those via find/replace).

    Returns ``{applied, added, unmatched, media_preserved}``. ``media_preserved=False``
    means an image changed and the caller MUST discard ``out_path``. Only the first
    occurrence of each ``find`` is replaced; a ``find`` not present → ``unmatched``.
    """
    from docx import Document

    before = _media_hashes(src_path)
    doc = Document(src_path)
    paras = list(_iter_paragraphs(doc))

    applied = 0
    unmatched: list[str] = []
    for c in corrections or []:
        find = str(c.get("find") or "")
        replace = str(c.get("replace") or "")
        if not find:
            continue
        hit = False
        for para in paras:
            if find in "".join(r.text for r in para.runs):
                if _replace_in_runs(para.runs, find, replace):
                    hit = True
                    applied += 1
                    break
        if not hit:
            unmatched.append(find)

    added = 0
    adds = [a for a in (additions or []) if str(a or "").strip()]
    if adds:
        doc.add_paragraph("")
        doc.add_paragraph("Reconciliation — requirements added back per the ratified plan:")
        for a in adds:
            doc.add_paragraph(f"• {a}")
            added += 1

    doc.save(out_path)
    media_preserved = (before == _media_hashes(out_path))
    return {"applied": applied, "added": added, "unmatched": unmatched, "media_preserved": media_preserved}
