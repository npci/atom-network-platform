# Copyright 2026 National Payments Corporation of India
# SPDX-License-Identifier: MIT

"""Styled Word (.docx) document builder — adapted from A2A-main's
docgen/tools/docx_builder.py, pruned to network platform needs.

Produces a single Document instance with:
  - Cover page (title, subtitle, doc_type, version, revision-history table)
  - Table of Contents (auto-refresh field)
  - Styled headings (Calibri body, color-coded H1/H2/H3)
  - Paragraphs, bullet lists, numbered lists
  - Styled tables (blue header row, alternating row colors)
  - Code blocks (Courier, grey background, left-indented)
  - Placeholder hook for diagrams at a target heading (feature deferred)

Usage:
    b = DocxBuilder()
    b.add_cover_page(title="BRD", subtitle="Crowd Funding", doc_type="BRD",
                     version="1.0", revision_history=[...])
    b.add_toc()
    b.add_heading("1. Executive Summary", level=1)
    b.add_paragraph("...")
    b.save("/path/out.docx")

The `add_section_content()` helper renders a GeneratedContent-shaped dict.
The docx_assembler module uses it to convert markdown → DOCX.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)


# ── Colour palette ──────────────────────────────────────────────────────────
BLUE_DARK = RGBColor(0x1D, 0x4E, 0xD8)   # header accents
NAVY      = RGBColor(0x0F, 0x2B, 0x6C)
DARK_GRAY = RGBColor(0x3B, 0x3B, 0x3B)
LIGHT_BLUE = "DCE7FA"                    # alternating row shade (hex string)
CODE_BG   = "F2F2F2"


# ── Inline markdown handling ─────────────────────────────────────────────────
# python-docx has no concept of markdown — paragraphs are just runs of plain
# text with formatting flags. The LLM-generated content uses CommonMark
# `**bold**` / `*italic*` / `` `code` `` markers throughout the body. Without
# this tokenizer those markers showed up literally in the .docx (the original
# bug: cover and first-page paragraphs displaying `**Feature**`).
#
# Two helpers:
#   _tokenize_inline_markdown(text) → [(segment, kind), ...] where kind is
#       one of "plain" / "bold" / "italic" / "bold_italic" / "code".
#   _strip_inline_markdown(text)   → plain string with markers removed.
#
# Both are deliberately conservative: they fix the patterns we've actually
# observed (orphan `***Word**`, paired `**Word**`, simple `*word*`, backtick
# code spans). They do not attempt to be a full CommonMark parser.

import re as _re

# Captures any of the inline marker types. The order matters — longest
# delimiters first so `***bold-italic***` doesn't get eaten as `**bold**`.
_INLINE_RE = _re.compile(
    r"\*\*\*([^*\n]+?)\*\*\*"        # ***bold italic***
    r"|\*\*([^*\n]+?)\*\*"            # **bold**
    r"|(?<!\*)\*([^*\s][^*\n]*?)\*(?!\*)"  # *italic* — avoid bullets
    r"|`([^`\n]+?)`"                  # `code`
)

# A code span on its own, used to re-scan the inside of an emphasis span.
_CODE_SPAN_RE = _re.compile(r"`([^`\n]+?)`")


def _emit_with_nested_code(out: list[tuple[str, str]], text: str, kind: str) -> None:
    """Append `text` as `kind`, splitting out any code span nested inside it.

    `_INLINE_RE` consumes an emphasis span in a single match, so backticks
    inside it were never re-scanned and rendered literally — ``**a `b` c**``
    came out as the visible text ``a `b` c`` in bold. Emphasis does not carry
    into the code fragment: it becomes a `code` run, the words either side keep
    `kind`. That keeps the segment vocabulary at the documented five values, so
    every existing consumer works unchanged.
    """
    if "`" not in text:
        out.append((text, kind))
        return
    cursor = 0
    for m in _CODE_SPAN_RE.finditer(text):
        if m.start() > cursor:
            out.append((text[cursor:m.start()], kind))
        out.append((m.group(1), "code"))
        cursor = m.end()
    if cursor < len(text):
        out.append((text[cursor:], kind))


def _tokenize_inline_markdown(text: str) -> list[tuple[str, str]]:
    """Split a paragraph string into (segment, kind) runs.

    Defensive cleanups before tokenizing:
      - `***Word**` (3 open + 2 close) → `**Word**`
      - `**Word***` (2 open + 3 close) → `**Word**`
      - Lone trailing `*` on an otherwise-balanced line → drop.
    """
    if not text:
        return []

    # Repair the malformed bold patterns the LLM most commonly emits before
    # feeding the regex below — otherwise asterisks end up in `plain` runs.
    text = _re.sub(r"\*{3}([^*\n]{1,80}?)\*{2}(?!\*)", r"**\1**", text)
    text = _re.sub(r"(?<!\*)\*{2}([^*\n]{1,80}?)\*{3}", r"**\1**", text)
    # Drop a stray trailing `*` if the line has an even count of single
    # asterisks (i.e. the trailing one is unpaired).
    if text.endswith("*") and not text.endswith("**"):
        singles = len(_re.findall(r"(?<!\*)\*(?!\*)", text))
        if singles % 2 == 1:
            text = text[:-1].rstrip()

    out: list[tuple[str, str]] = []
    cursor = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > cursor:
            out.append((text[cursor:m.start()], "plain"))
        bold_italic, bold, italic, code = m.group(1), m.group(2), m.group(3), m.group(4)
        if bold_italic is not None:
            _emit_with_nested_code(out, bold_italic, "bold_italic")
        elif bold is not None:
            _emit_with_nested_code(out, bold, "bold")
        elif italic is not None:
            _emit_with_nested_code(out, italic, "italic")
        elif code is not None:
            out.append((code, "code"))
        cursor = m.end()
    if cursor < len(text):
        out.append((text[cursor:], "plain"))
    # Drop empty segments — they happen when a line starts with markers.
    return [(s, k) for s, k in out if s]


def _strip_inline_markdown(text: str) -> str:
    """Remove inline markdown markers, keeping only the inner text. Used for
    headings where Word's own heading style supplies the visual emphasis."""
    if not text:
        return text or ""
    return "".join(seg for seg, _ in _tokenize_inline_markdown(text))


# ──────────────────────────────────────────────────────────────────────────────
# DocxBuilder
# ──────────────────────────────────────────────────────────────────────────────

class DocxBuilder:
    """Wraps python-docx with the Authority styling + higher-level helpers."""

    def __init__(self):
        self.doc = Document()
        self._setup_default_style()
        self._setup_page_size()

    # ── Setup ──────────────────────────────────────────────────────────────

    def _setup_default_style(self) -> None:
        style = self.doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        # Heading styles
        for lvl, color in ((1, BLUE_DARK), (2, NAVY), (3, DARK_GRAY)):
            try:
                h = self.doc.styles[f"Heading {lvl}"]
                h.font.name = "Calibri"
                h.font.size = Pt({1: 18, 2: 14, 3: 12}[lvl])
                h.font.bold = True
                h.font.color.rgb = color
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(6)
            except KeyError:
                pass

    def _setup_page_size(self) -> None:
        for section in self.doc.sections:
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
            section.top_margin = section.bottom_margin = Inches(1)
            section.left_margin = section.right_margin = Inches(1)

    # ── Cover page ─────────────────────────────────────────────────────────

    def add_cover_page(
        self,
        *,
        title: str,
        subtitle: str = "",
        doc_type: str = "",
        version: str = "1.0",
        revision_history: list[dict] | None = None,
    ) -> None:
        # Title — strip any inline markdown markers; the cover sets its own
        # bold styling, so leftover `**` would just look broken.
        clean_title = _strip_inline_markdown(title)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(120)
        r = p.add_run(clean_title)
        r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = NAVY

        # Subtitle
        if subtitle:
            clean_subtitle = _strip_inline_markdown(subtitle)
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(clean_subtitle)
            r.font.size = Pt(16); r.font.color.rgb = DARK_GRAY

        # Doc type + version
        meta = self.doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.paragraph_format.space_before = Pt(60)
        meta_run = meta.add_run(f"{doc_type}    •    Version {version}")
        meta_run.font.size = Pt(12); meta_run.font.color.rgb = DARK_GRAY

        # Revision history table
        if revision_history:
            self.doc.add_paragraph().paragraph_format.space_before = Pt(80)
            hdr = self.doc.add_paragraph("Revision History")
            hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in hdr.runs:
                r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = NAVY

            headers = ["Version", "Date", "Author", "Changes"]
            data = [[
                str(row.get("version", "")),
                str(row.get("date", "")),
                str(row.get("author", "")),
                str(row.get("changes", "")),
            ] for row in revision_history]
            self.add_styled_table(headers=headers, rows=data)

        self.doc.add_page_break()

    # ── Table of Contents ──────────────────────────────────────────────────

    def add_toc(self) -> None:
        """Insert a TOC field that Word auto-populates when opened."""
        p = self.doc.add_paragraph()
        heading_run = p.add_run("Table of Contents")
        heading_run.font.size = Pt(18); heading_run.font.bold = True
        heading_run.font.color.rgb = NAVY

        # Field code for TOC
        p = self.doc.add_paragraph()
        run = p.add_run()
        fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
        instr.text = 'TOC \\o "1-3" \\h \\z \\u'
        fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        # Placeholder text before Word refreshes
        placeholder = OxmlElement("w:t"); placeholder.text = "Right-click → Update Field to populate."
        run._r.append(placeholder)
        run._r.append(fld_end)

        self.doc.add_page_break()

    # ── Headings / paragraphs / lists ──────────────────────────────────────

    def add_heading(self, text: str, level: int = 1) -> None:
        # Strip inline emphasis markers from headings — Word headings carry
        # their own bold styling, so leftover `**…**` would just look ugly.
        clean = _strip_inline_markdown(text)
        self.doc.add_heading(clean, level=max(1, min(level, 3)))

    def add_paragraph(self, text: str) -> None:
        if not text or not text.strip():
            return
        # WHY: the markdown→docx assembler used to call this with raw text
        # like "Introduce **the network Corporate Circle**, a NET-based ecosystem…"
        # python-docx doesn't interpret markdown, so the `**` showed up
        # literally on the first page (and everywhere else). We now split
        # the text into runs at `**bold**` and `*italic*` boundaries and
        # apply the right run.font flags. Falls back to a plain run for
        # any segment without markers.
        p = self.doc.add_paragraph()
        for segment, kind in _tokenize_inline_markdown(text.strip()):
            run = p.add_run(segment)
            if kind == "bold":
                run.font.bold = True
            elif kind == "italic":
                run.font.italic = True
            elif kind == "bold_italic":
                run.font.bold = True
                run.font.italic = True
            elif kind == "code":
                run.font.name = "Courier New"

    def add_bullets(self, items: list[str]) -> None:
        for item in items:
            s = str(item).strip()
            if not s:
                continue
            p = self.doc.add_paragraph(style="List Bullet")
            for segment, kind in _tokenize_inline_markdown(s):
                run = p.add_run(segment)
                if kind == "bold":
                    run.font.bold = True
                elif kind == "italic":
                    run.font.italic = True
                elif kind == "bold_italic":
                    run.font.bold = True
                    run.font.italic = True
                elif kind == "code":
                    run.font.name = "Courier New"

    def add_numbered(self, items: list[str]) -> None:
        for item in items:
            s = str(item).strip()
            if not s:
                continue
            p = self.doc.add_paragraph(style="List Number")
            for segment, kind in _tokenize_inline_markdown(s):
                run = p.add_run(segment)
                if kind == "bold":
                    run.font.bold = True
                elif kind == "italic":
                    run.font.italic = True
                elif kind == "bold_italic":
                    run.font.bold = True
                    run.font.italic = True
                elif kind == "code":
                    run.font.name = "Courier New"

    # ── Tables ──────────────────────────────────────────────────────────────

    def add_styled_table(
        self,
        headers: list[str],
        rows: list[list[str]],
        *,
        caption: str = "",
    ) -> None:
        if not headers:
            return
        if caption:
            p = self.doc.add_paragraph(caption)
            for r in p.runs:
                r.font.italic = True; r.font.size = Pt(10); r.font.color.rgb = DARK_GRAY

        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        # Header row
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = str(h)
            self._shade_cell(cell, "1D4ED8")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.size = Pt(11)
        # Body rows
        for ri, row in enumerate(rows, start=1):
            for ci in range(len(headers)):
                val = row[ci] if ci < len(row) else ""
                cell = table.rows[ri].cells[ci]
                cell.text = str(val)
                if ri % 2 == 0:
                    self._shade_cell(cell, LIGHT_BLUE)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Space after
        self.doc.add_paragraph()

    def _shade_cell(self, cell, hex_color: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)

    # ── Code blocks ────────────────────────────────────────────────────────

    def add_code_block(self, text: str, language: str = "") -> None:
        if not text or not text.strip():
            return
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = "Courier New"; run.font.size = Pt(9)
        run.font.color.rgb = DARK_GRAY
        # Shade paragraph background
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), CODE_BG)
        pPr.append(shd)

    # ── Diagram hook (deferred feature — placeholder caption only) ────────

    def add_diagram_placeholder(self, description: str) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[Diagram: {description}]")
        run.font.italic = True; run.font.color.rgb = DARK_GRAY

    # ── Save ──────────────────────────────────────────────────────────────

    def save(self, path: str | Path) -> Path:
        out = Path(path)
        out.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(out))
        logger.info("DocxBuilder saved: %s", out)
        return out


# ──────────────────────────────────────────────────────────────────────────────
# Circular-specific variant (Times New Roman letterhead + the Authority footer).
# ──────────────────────────────────────────────────────────────────────────────

class CircularDocxBuilder(DocxBuilder):
    """Specialised builder for the Authority Circulars — letterhead, narrow margins, serif font."""

    def _setup_default_style(self) -> None:
        style = self.doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        for lvl in (1, 2, 3):
            try:
                h = self.doc.styles[f"Heading {lvl}"]
                h.font.name = "Times New Roman"
                h.font.bold = True
                h.font.color.rgb = DARK_GRAY
            except KeyError:
                pass

    def _setup_page_size(self) -> None:
        for section in self.doc.sections:
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.5)

    def add_letterhead(self, reference: str = "", date: str = "") -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("NATIONAL PAYMENTS CORPORATION OF INDIA")
        r.font.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run("Unified Payments Interface (the network)")
        r2.font.italic = True; r2.font.size = Pt(11); r2.font.color.rgb = DARK_GRAY

        # Horizontal line by drawing a paragraph with bottom border
        self.doc.add_paragraph("─" * 80)

        # Reference + date block
        if reference or date:
            p = self.doc.add_paragraph()
            if reference:
                r = p.add_run(f"Ref: {reference}")
                r.font.bold = True
            if date:
                p.add_run("\t" * 4)
                rd = p.add_run(f"Date: {date}")
                rd.font.bold = True

    def add_npci_footer(self) -> None:
        self.doc.add_paragraph("─" * 80)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Operator-supplied. The organisation's postal address and contact
        # details are deployment identity, not platform code — a fork must not
        # emit the Authority's address on its documents. Set DOC_FOOTER_TEXT; the footer
        # rule is still drawn when it is empty, so layout is unchanged.
        from app.core.config import settings
        footer_text = (getattr(settings, "doc_footer_text", "") or "").strip()
        if not footer_text:
            return
        r = p.add_run(footer_text)
        r.font.size = Pt(9); r.font.color.rgb = DARK_GRAY
