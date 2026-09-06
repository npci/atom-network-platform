# Copyright 2026 National Payments Corporation of India
# SPDX-License-Identifier: MIT

"""Render the NPCI Certification Result (cert sign-off) certificate as a .docx.

Reproduces the layout of the official certificate (see
`docs/cert_signoff_template.pdf`):
a salutation, a header table of certification fields with tick-box lines, a
per-test-case results table, and the standard Note + Disclaimer blocks.

We generate from scratch with python-docx rather than substituting into a
template because (a) we don't have the editable .docx, and (b) the results
table has a variable number of rows (one per test case) that no fixed
template could carry.

`build_signoff_docx(meta, results) -> bytes` returns the file in memory so the
caller can base64 it onto the `cert_completion_signoff` A2A task without a
disk round-trip.

Checkbox fields are passed as the *selected* label(s); the renderer holds the
full option list per field (matching the certificate) and ticks the matches.
Fields the platform does not store fall back to STATIC_DEFAULTS, which is the
one place to edit the boilerplate version strings.
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

CHECKED = "☑"    # ☑
UNCHECKED = "☐"  # ☐

# ── Pack-driven vocabulary (genericisation sweep) ────────────────────────────
# The certificate's option lists and boilerplate are the certification BODY's
# words, not the engine's: scope options come from the pack's cert-vocabulary
# role scopes, the product list / version strings / member-type options from
# its prompt blocks, and every authority mention from the `authority` block.
# For UPI the rendered bytes are identical to the previous hardcoded layout.
from app.core.domain.contract import cert_vocabulary_of
from app.core.domain.registry import get_active_pack, prompt_block

_AUTHORITY = prompt_block("authority", "the certifying authority")
_AUTHORITY_FULL = prompt_block("authority_full_name", "") or _AUTHORITY
_PARTNER = prompt_block("cert_partner_label", "Bank")


def _block_options(name: str) -> list[str]:
    """A `|`-separated option list from a pack prompt block; [] when absent —
    the tick-box line then renders empty, which is a true statement about a
    domain whose certificates don't carry that field."""
    raw = prompt_block(name, "")
    return [p.strip() for p in raw.split("|") if p.strip()] if raw else []


def _scope_options() -> list[str]:
    """Distinct certification scopes in declaration order, plus the combined
    option when the domain has more than one (UPI: ACQUIRER / ISSUER /
    "ACQUIRER +ISSUER" — spacing preserved from the official certificate)."""
    scopes: list[str] = []
    for role_scopes in cert_vocabulary_of(get_active_pack()).role_scopes.values():
        for s in role_scopes:
            if s not in scopes:
                scopes.append(s)
    if len(scopes) > 1:
        scopes.append(" +".join(scopes))
    return scopes


def _product_options() -> list[str]:
    """The certificate's product tick-list: the pack's
    `cert_certified_product_options` block when declared (UPI lists products
    beyond its `product_labels` mapping), else the distinct product labels."""
    declared = _block_options("cert_certified_product_options")
    if declared:
        return declared
    out: list[str] = []
    for _kw, label in cert_vocabulary_of(get_active_pack()).product_labels:
        if label not in out:
            out.append(label)
    return out


# Option lists per tick-box field, in the order they appear on the certificate.
_OPTIONS: dict[str, list[str]] = {
    "type_of_bank":   _block_options("cert_type_of_bank_options"),
    "certified_product": _product_options(),
    "scope":          _scope_options(),
    "channel":        ["MOBILE", "WEBCOLLECT", "USSD"],
    "os":             ["ANDROID", "IOS"],
    "transaction_messaging": ["XML – API"],
    "risk_team_approval":    ["YES", "NO", f"{_AUTHORITY} Business to confirm"],
    "final_result":          ["PASSED", "FAILED"],
    "online_certification":  ["PASSED", "FAILED", "WAIVERED"],
    "offline_certification": ["PASSED", "FAILED", "WAIVERED", "NOT APPLICABLE"],
    "application_certification": ["PASSED", "FAILED", "WAIVERED", "NOT APPLICABLE"],
}

# Boilerplate for fields the platform doesn't store. The version strings come
# from the pack; edit the pack, not this dict.
STATIC_DEFAULTS: dict[str, object] = {
    "org_id":                "",
    "type_of_bank":          [],           # ops ticks in the downloaded doc
    "channel":               [],
    "os":                    [],
    "service_provider":      prompt_block("cert_service_provider", ""),
    "spec_version":          prompt_block("cert_spec_version", ""),
    "testcase_version":      "",
    "transaction_messaging": ["XML – API"],
    "risk_team_approval":    [f"{_AUTHORITY} Business to confirm"],
    "offline_certification":     ["NOT APPLICABLE"],
    "application_certification": ["NOT APPLICABLE"],
}

_NOTE_LINES = [
    f"Certified the above mentioned {_PARTNER} to use the mentioned features for their "
    f"customers as the {_PARTNER} has successfully completed certification with {_AUTHORITY} as "
    "above highlighted date. This certificate is valid up to 50 days from certified date.",
    f"The certified {_PARTNER} to move into production within 3 months of being certified.",
    f"Sign the document with {_PARTNER} stamp on each page and submit the signed copy "
    f"within 7 days to {_AUTHORITY}.",
]

_DISCLAIMER = (
    "The information contained in this certificate and any attachments to this "
    "certificate are confidential and company privileged information. This "
    f"certificate is only used for proving the host system compliance with {_AUTHORITY} "
    "and this shall not be distributed or disclosed to any other third parties "
    f"without the written permission from {_AUTHORITY}."
)

_SALUTATION = (
    f"Sir / Madam,\nWe ({_AUTHORITY_FULL}) are glad to inform "
    "you that, all the required test cases have been completed with results as "
    "shown in the following table:"
)


def _as_set(value: object) -> set[str]:
    """Normalise a selected-label value (str | iterable | None) to a lowercased set."""
    if value is None:
        return set()
    if isinstance(value, str):
        items = [value]
    else:
        items = list(value)
    return {str(v).strip().lower() for v in items if str(v).strip()}


def _checkbox_text(field: str, selected: object) -> str:
    """Render a tick-box line for `field`, ticking labels in `selected`."""
    chosen = _as_set(selected)
    parts = []
    for label in _OPTIONS[field]:
        mark = CHECKED if label.lower() in chosen else UNCHECKED
        parts.append(f"{mark} {label}")
    return "    ".join(parts)


def _add_kv_row(table, label: str, value: str) -> None:
    row = table.add_row().cells
    row[0].text = label
    row[1].text = value or ""
    row[0].paragraphs[0].runs and setattr(row[0].paragraphs[0].runs[0].font, "bold", True)


def build_signoff_docx(meta: dict, results: list[dict]) -> bytes:
    """Build the certification sign-off .docx and return its bytes.

    `meta` carries the header fields; missing keys fall back to STATIC_DEFAULTS.
    `results` is a list of {test_id, txn_id, date, status} (one per test case).
    """
    m = {**STATIC_DEFAULTS, **(meta or {})}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Salutation
    for line in _SALUTATION.split("\n"):
        doc.add_paragraph(line)

    # Header table — title row spanning both columns, then label/value rows.
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    title_cells = table.rows[0].cells
    title_cells[0].merge(title_cells[1])
    tpara = title_cells[0].paragraphs[0]
    tpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = tpara.add_run(f"{_AUTHORITY} Certification – Result")
    trun.bold = True

    _add_kv_row(table, "Org ID", str(m.get("org_id", "")))
    _add_kv_row(table, "HANDLE", str(m.get("handle", "")))
    _add_kv_row(table, "Bank Name", str(m.get("bank_name", "")))
    _add_kv_row(table, "Certification Id", str(m.get("certification_id", "")))
    _add_kv_row(table, "Type of Bank", _checkbox_text("type_of_bank", m.get("type_of_bank")))
    _add_kv_row(table, "Date of Certification", str(m.get("date_of_certification", "")))
    _add_kv_row(table, "Certified Product", _checkbox_text("certified_product", m.get("certified_product")))
    _add_kv_row(table, "Scope of Certification", _checkbox_text("scope", m.get("scope")))
    _add_kv_row(table, "Channel", _checkbox_text("channel", m.get("channel")))
    _add_kv_row(table, "OS", _checkbox_text("os", m.get("os")))
    _add_kv_row(table, "Service Provider", str(m.get("service_provider", "")))
    _add_kv_row(table, "Specification version", str(m.get("spec_version", "")))
    _add_kv_row(table, "Test case version", str(m.get("testcase_version", "")))
    _add_kv_row(table, "Rounds of testing", str(m.get("rounds", "")))
    _add_kv_row(table, "Script documented date", str(m.get("script_documented_date", m.get("date_of_certification", ""))))
    _add_kv_row(table, "Transaction messaging", _checkbox_text("transaction_messaging", m.get("transaction_messaging")))
    _add_kv_row(table, f"Approval from {_AUTHORITY} Risk Team", _checkbox_text("risk_team_approval", m.get("risk_team_approval")))

    # Results row — a single value cell holding the four certification verdicts.
    res_cells = table.add_row().cells
    res_cells[0].text = "Results"
    if res_cells[0].paragraphs[0].runs:
        res_cells[0].paragraphs[0].runs[0].font.bold = True
    rc = res_cells[1]
    rc.paragraphs[0].text = f"Final Result: {_checkbox_text('final_result', m.get('final_result'))}"
    rc.add_paragraph(f"Online Certification: {_checkbox_text('online_certification', m.get('online_certification'))}")
    rc.add_paragraph(f"Offline Certification: {_checkbox_text('offline_certification', m.get('offline_certification'))}")
    rc.add_paragraph(f"Application Certification: {_checkbox_text('application_certification', m.get('application_certification'))}")

    # Note + Disclaimer
    doc.add_paragraph()
    note_head = doc.add_paragraph()
    note_head.add_run("Note:").bold = True
    for i, line in enumerate(_NOTE_LINES, start=1):
        doc.add_paragraph(f"{i}. {line}")
    disc_head = doc.add_paragraph()
    disc_head.add_run("Disclaimer:").bold = True
    doc.add_paragraph(_DISCLAIMER)

    # Per-test-case results table
    doc.add_paragraph()
    rtable = doc.add_table(rows=1, cols=4)
    rtable.style = "Table Grid"
    hdr = rtable.rows[0].cells
    for cell, label in zip(hdr, ("TEST ID", "TXN ID", "Date", "Status")):
        cell.paragraphs[0].add_run(label).bold = True
    for r in results or []:
        status_raw = str(r.get("status", "")).upper()
        status = "Success" if status_raw in ("PASS", "PASSED", "SUCCESS") else (
            "Failed" if status_raw in ("FAIL", "FAILED") else r.get("status", "")
        )
        cells = rtable.add_row().cells
        cells[0].text = str(r.get("test_id") or r.get("test_case_id") or "")
        cells[1].text = str(r.get("txn_id") or r.get("correlation_id") or "")
        cells[2].text = str(r.get("date") or "")
        cells[3].text = str(status)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# Sample values mirroring `uat script.pdf` — used by the smoke test below and
# handy as a reference for the meta shape the orchestrator must produce.
_SAMPLE_META = {
    "org_id": "159049",
    "handle": "fedepsp",
    "bank_name": "Federal Bank",
    "certification_id": "NP2FDRLUPI240840318",
    "type_of_bank": ["Direct Member"],
    "date_of_certification": "22 August 2024",
    "certified_product": ["UPI -Lite Autopay Issuer"],
    "scope": [],
    "channel": ["MOBILE"],
    "os": ["ANDROID", "IOS"],
    "service_provider": "Bank Inhouse development",
    "spec_version": "Unified Payments Interface 2.0 TSD v1.42",
    "testcase_version": "UPI Lite Autopay Testcases V5.0",
    "rounds": "2",
    "script_documented_date": "22 August 2024",
    "transaction_messaging": ["XML – API"],
    "risk_team_approval": ["NPCI Business to confirm"],
    "final_result": ["PASSED"],
    "online_certification": ["PASSED"],
}

# Demo rows for the `__main__` block below — nothing imports these. The txn ids
# are GENERATED rather than written out as literals: the previous fixtures were
# real-looking 32-hex NPCI transaction references copied from a certification
# run, which is both a secret-scanner magnet (SCR #7 — TruffleHog HighEntropy
# Strings) and a bad habit, since nothing distinguishes a pasted sample from a
# pasted production value at review time. A deterministic seed keeps the sample
# document byte-stable across runs, and the SAMPLE_ prefix makes the origin of
# any id that escapes into a document self-evident.
_SAMPLE_RESULTS = [
    {
        "test_id": f"TC_{i:02d}",
        # Not a secret and not random: a fixed, obviously-synthetic pattern that
        # keeps the real format's shape (LIT prefix + 32 chars) so the generated
        # document still lays out realistically.
        "txn_id": f"LITSAMPLE{i:023d}",
        "date": "22-08-2024",
        "status": "PASS",
    }
    for i in range(1, 4)
]


if __name__ == "__main__":
    data = build_signoff_docx(_SAMPLE_META, _SAMPLE_RESULTS)
    with open("cert_signoff_sample.docx", "wb") as f:
        f.write(data)
    print(f"wrote cert_signoff_sample.docx ({len(data)} bytes)")
