# Copyright 2026 National Payments Corporation of India
# SPDX-License-Identifier: MIT

"""Render the kit "Summary of Changes" as a .docx.

The version change summary (plan.summary / the envelope's change_summary) is the
partner-facing note describing what changed from v(N) to v(N+1). This builds it
as a downloadable Word document so it can sit alongside the kit documents.
"""
from __future__ import annotations

import io

from docx import Document
from docx.shared import Pt


def build_change_summary_docx(*, change_title: str, version: int, summary: str) -> bytes:
    """Return a .docx (bytes) of the change summary for the given version."""
    doc = Document()

    title = doc.add_heading("Summary of Changes", level=0)
    sub = doc.add_paragraph()
    run = sub.add_run(f"{change_title or 'Change'} — Product Kit v{version}")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()  # spacer

    # The summary is plain prose; render blank-line-separated blocks as
    # paragraphs and single newlines as line breaks within a paragraph.
    for block in (summary or "").split("\n\n"):
        block = block.strip()
        if not block:
            continue
        p = doc.add_paragraph()
        for i, line in enumerate(block.split("\n")):
            if i:
                p.add_run().add_break()
            p.add_run(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
