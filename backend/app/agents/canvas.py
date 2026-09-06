# Copyright 2026 National Payments Corporation of India
# SPDX-License-Identifier: MIT

"""Product Canvas Generator agent.

Produces a structured Product Canvas matching the "Build Framework" template
(10 sections) from the enriched prompt and deep-research report, then handles
iterative feedback refinement.

Also provides generate_canvas_docx() to export the canvas as a .docx file
that mirrors the template's grid layout.
"""
import io
from app.core.prompts import render_prompt
import logging
import re
from collections.abc import AsyncGenerator

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

from app.core.llm import stream_llm
from app.core.domain.registry import prompt_block

# Supplied by the active domain pack, not imported from a UPI module.
NETWORK_HARD_RULES = prompt_block("hard_rules")
from app.agents.blueprints import format_for_prompt as _blueprint_for_prompt
from app.agents._prompt_safety import wrap_untrusted, ANTI_INJECTION_CLAUSE

logger = logging.getLogger(__name__)

_CANVAS_BLUEPRINT_BLOCK = _blueprint_for_prompt("canvas")

SYSTEM_PROMPT = render_prompt(
    "agents/canvas/system_prompt.md",
    NETWORK_HARD_RULES=NETWORK_HARD_RULES, CANVAS_BLUEPRINT_BLOCK=_CANVAS_BLUEPRINT_BLOCK, ANTI_INJECTION_CLAUSE=ANTI_INJECTION_CLAUSE,
    PLATFORM_NAME=prompt_block("platform_name", "this change-management platform"),
    ECOSYSTEM_ACTORS=prompt_block("ecosystem_actors", "the platform's ecosystem"),
    REGULATORY_BODY=prompt_block("regulatory_body", "the applicable regulator"),
    PRODUCT_OPERATING_EXTRA=prompt_block("product_operating_extra", ""),
    AUTHORITY=prompt_block("authority", "the platform operator"),
    REFERENCE_KIND=prompt_block("reference_kind", "published guidance"),
    DOMAIN_NAME=prompt_block("domain_name", "this domain"),
)


async def stream_canvas_turn(
    enriched_prompt: str,
    research_report: str,
    conversation_history: list[dict],
    new_user_message: str,
) -> AsyncGenerator[str, None]:
    """Stream a canvas generation or refinement turn."""
    context = f"""ENRICHED PROMPT:
{wrap_untrusted(enriched_prompt, "ENRICHED_PROMPT")}

---
RESEARCH REPORT:
{wrap_untrusted(research_report, "RESEARCH_REPORT")}
"""
    messages = conversation_history + [{"role": "user", "content": wrap_untrusted(new_user_message, "USER_MESSAGE")}]
    if len(conversation_history) == 0:
        messages = [{"role": "user", "content": f"{context}\n\n---\n{wrap_untrusted(new_user_message, 'USER_MESSAGE')}"}]

    logger.info("CanvasAgent — streaming turn, history_len=%d", len(messages))

    # Slice 9b — preserve [N] citations from upstream research_report.
    from app.core.config import settings as _settings
    from app.agents.citations import preserve_suffix, upstream_has_citations
    system_prompt = SYSTEM_PROMPT
    if _settings.use_citation_enforcement and upstream_has_citations(research_report):
        system_prompt = SYSTEM_PROMPT + preserve_suffix()

    # max_tokens bumped 3000 → 24000 (2026-05-04, Layer-3 of truncation fix).
    # User reported visible truncation in product canvas. The 3000-token cap
    # was extremely tight for Claude Sonnet 4.6 — production observed 33k+
    # char canvas outputs (~8k tokens) cut mid-section. Sonnet supports 64K
    # output, so 24K is generous for any plausible canvas length.
    async for chunk in stream_llm(system=system_prompt, messages=messages, max_tokens=24000, agent_name="canvas"):
        yield chunk


# ── DOCX export ───────────────────────────────────────────────────────────────

def _parse_sections(content: str) -> dict[str, str]:
    """Parse the 10 canvas sections from markdown content."""
    section_pattern = re.compile(r'^## (\d+)\. (.+)$', re.MULTILINE)
    matches = list(section_pattern.finditer(content))
    sections = {}
    for i, m in enumerate(matches):
        key = f"{m.group(1)}. {m.group(2).strip()}"
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(content)
        sections[key] = content[start:end].strip()
    return sections


def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _add_cell_content(cell, title: str, body: str, title_color='1A3C5E'):
    """Write bold title + body text into a table cell."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    para = cell.paragraphs[0]
    para.clear()

    # Title
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(
        int(title_color[0:2], 16),
        int(title_color[2:4], 16),
        int(title_color[4:6], 16),
    )

    # Body lines
    for line in body.split('\n'):
        line = line.strip()
        if not line:
            continue
        # Strip markdown bold/italic markers
        line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
        line = re.sub(r'\*(.+?)\*', r'\1', line)
        # Strip leading bullet chars
        line = re.sub(r'^[-•]\s*', '', line)

        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(1)

        # Sub-label detection: "Label — rest" or "Label:"
        if re.match(r'^[A-Z].+?[:—–-]', line):
            parts = re.split(r'[:—–-]', line, 1)
            label_run = p.add_run(parts[0].strip() + ': ')
            label_run.bold = True
            label_run.font.size = Pt(8)
            label_run.font.color.rgb = RGBColor(80, 80, 80)
            if len(parts) > 1:
                rest_run = p.add_run(parts[1].strip())
                rest_run.font.size = Pt(8)
        else:
            r = p.add_run(line)
            r.font.size = Pt(8)


def generate_canvas_docx(change_title: str, canvas_content: str) -> bytes:
    """
    Generate a .docx Product Canvas that mirrors the NPCI Build Framework grid.

    Layout (matching the PDF template):
    ┌─────────────────────────────────────────────────────────┐
    │  Header: "Build framework for <title>"                  │
    ├─────────────────────────────────────────────────────────┤
    │  1. Feature  (full width)                               │
    ├──────────────────┬──────────────────┬───────────────────┤
    │  2. Need         │  3. Market View  │  4. Scalability   │
    ├────────┬─────────┴──────────────────┤  (cont.)          │
    │  5.    │  6. Product Operating      │                   │
    │  Valid ├──────────────────────────────────────────────── │
    │  ation │  7. Product Comms                              │
    ├────────┴──────┬──────────────────┬─────────────────────┤
    │  8. Pricing   │  9. Risks        │  10. Compliance     │
    └───────────────┴──────────────────┴─────────────────────┘
    """
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    sections = _parse_sections(canvas_content)

    def get(num: int) -> tuple[str, str]:
        for k, v in sections.items():
            if k.startswith(f"{num}."):
                return k, v
        return f"{num}.", ""

    # ── Header ────────────────────────────────────────────────
    hdr = doc.add_paragraph()
    hdr.paragraph_format.space_after = Pt(6)
    run = hdr.add_run(f"Build framework for {change_title}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(26, 60, 94)  # NPCI navy

    # ── Row 1: Feature (full width) ───────────────────────────
    t1 = doc.add_table(rows=1, cols=1)
    t1.style = 'Table Grid'
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t1.columns[0].width = Inches(7.5)
    k1, v1 = get(1)
    _set_cell_bg(t1.cell(0, 0), 'EBF3FB')
    _add_cell_content(t1.cell(0, 0), '1. Feature', v1)
    doc.add_paragraph()

    # ── Row 2: Need | Market View | Scalability ───────────────
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = 'Table Grid'
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths2 = [Inches(2.8), Inches(2.8), Inches(1.9)]
    for i, w in enumerate(widths2):
        t2.columns[i].width = w

    _, v2 = get(2)
    _, v3 = get(3)
    _, v4 = get(4)
    _set_cell_bg(t2.cell(0, 0), 'FFFFFF')
    _set_cell_bg(t2.cell(0, 1), 'FFFFFF')
    _set_cell_bg(t2.cell(0, 2), 'FFFFFF')
    _add_cell_content(t2.cell(0, 0), '2. Need', v2)
    _add_cell_content(t2.cell(0, 1), '3. Market View', v3)
    _add_cell_content(t2.cell(0, 2), '4. Scalability', v4)
    doc.add_paragraph()

    # ── Row 3: Validation | Product Operating | Product Comms ─
    t3 = doc.add_table(rows=1, cols=3)
    t3.style = 'Table Grid'
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths3 = [Inches(1.9), Inches(2.8), Inches(2.8)]
    for i, w in enumerate(widths3):
        t3.columns[i].width = w

    _, v5 = get(5)
    _, v6 = get(6)
    _, v7 = get(7)
    _set_cell_bg(t3.cell(0, 0), 'FFFFFF')
    _set_cell_bg(t3.cell(0, 1), 'FFFFFF')
    _set_cell_bg(t3.cell(0, 2), 'FFFFFF')
    _add_cell_content(t3.cell(0, 0), '5. Validation', v5)
    _add_cell_content(t3.cell(0, 1), '6. Product Operating', v6)
    _add_cell_content(t3.cell(0, 2), '7. Product Comms\n(external + internal)', v7)
    doc.add_paragraph()

    # ── Row 4: Pricing | Risks | Compliance ───────────────────
    t4 = doc.add_table(rows=1, cols=3)
    t4.style = 'Table Grid'
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths4 = [Inches(1.9), Inches(2.8), Inches(2.8)]
    for i, w in enumerate(widths4):
        t4.columns[i].width = w

    _, v8  = get(8)
    _, v9  = get(9)
    _, v10 = get(10)
    _set_cell_bg(t4.cell(0, 0), 'FFFFFF')
    _set_cell_bg(t4.cell(0, 1), 'FFFFFF')
    _set_cell_bg(t4.cell(0, 2), 'FFFFFF')
    _add_cell_content(t4.cell(0, 0), '8. Pricing', v8)
    _add_cell_content(t4.cell(0, 1), '9. Potential Risks', v9)
    _add_cell_content(t4.cell(0, 2), '10. Compliance', v10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
